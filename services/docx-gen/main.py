import os
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional, List, Dict
import datetime
import tempfile
import io

app = FastAPI(title="LicitAI DOCX Generator")

class DocumentRequest(BaseModel):
    convocante: str
    licitacion_no: str
    objeto: Optional[str] = None
    empresa_nombre: str
    empresa_rfc: str
    representante_legal: str
    cargo_representante: str
    tipo_persona: str  # 'PF' o 'PM'
    titulo_documento: str
    contenido: List[Dict[str, str]]  # [{'tipo': 'parrafo', 'texto': '...'}, {'tipo': 'tabla', 'headers': [], 'data': []}]
    show_header: bool = True
    show_footer: bool = True
    logo_path: Optional[str] = None
    domicilio_fiscal: Optional[str] = None

def add_header(doc, convocante, licitacion_no, logo_path=None):
    section = doc.sections[0]
    header = section.header
    
    # Limpiar párrafos existentes en el header
    for p in header.paragraphs:
        p.text = ""

    # Usamos una tabla invisible para alinear logo (izquierda) y datos (derecha)
    htable = header.add_table(1, 2, width=Inches(6.5))
    htable.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Celda 0: Logo
    img_cell = htable.rows[0].cells[0]
    if logo_path and os.path.exists(logo_path):
        try:
            p_img = img_cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_img = p_img.add_run()
            run_img.add_picture(logo_path, width=Inches(1.4))
        except Exception as e:
            img_cell.text = "[LOGO]"
    else:
        img_cell.text = ""

    # Celda 1: Datos Licitación
    txt_cell = htable.rows[0].cells[1]
    p_txt = txt_cell.paragraphs[0]
    p_txt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_conv = p_txt.add_run(f"{convocante.upper()}\n")
    run_conv.font.bold = True
    run_conv.font.size = Pt(11)
    
    run_lic = p_txt.add_run(f"LICITACIÓN: {licitacion_no}")
    run_lic.font.bold = True
    run_lic.font.color.rgb = None # Negro
    run_lic.font.size = Pt(10)

def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento generado por LicitAI-LM | Inteligencia Artificial para Licitaciones")
    run.font.size = Pt(8)
    run.font.italic = True

@app.post("/docx/generate")
async def generate_document(req: DocumentRequest):
    try:
        doc = Document()
        
        # Estilos globales (Arial 11pt, Justificado)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Márgenes (2.5 cm = 0.98 inches)
        for section in doc.sections:
            section.top_margin = Inches(0.98)
            section.bottom_margin = Inches(0.98)
            section.left_margin = Inches(0.98)
            section.right_margin = Inches(0.98)

        # 1. Header
        if req.show_header:
            add_header(doc, req.convocante, req.licitacion_no, req.logo_path)
            
        # 2. Title
        doc.add_paragraph("\n")
        title = doc.add_heading(req.titulo_documento, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2.1 Datos de la Empresa (RFC, Domicilio, Tipo de Persona)
        info_block = []
        today = datetime.date.today().strftime("%d/%m/%Y")
        info_block.append(f"Fecha: {today}")
        if req.empresa_nombre:
            info_block.append(f"Licitante: {req.empresa_nombre.upper()}")
        if req.empresa_rfc:
            info_block.append(f"R.F.C.: {req.empresa_rfc}")
        
        if info_block:
            p_info = doc.add_paragraph("\n".join(info_block))
            p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p_info.runs:
                run.font.size = Pt(9)
        
        doc.add_paragraph("\n")
        
        intro_dom = req.domicilio_fiscal or "[DATOS POR COMPLETAR]"
        intro_rfc = req.empresa_rfc or "[DATOS POR COMPLETAR]"
        intro_obj = req.objeto or "el suministro de bienes/servicios"
        
        intro_line = f"En mi carácter de representante legal de la empresa {req.empresa_nombre.upper()}, con domicilio fiscal en {intro_dom} y R.F.C. {intro_rfc}, me dirijo a la {req.convocante.upper()}, con objeto de presentar nuestra propuesta para la licitación {req.licitacion_no}, relativo a {intro_obj}."
        
        p_intro = doc.add_paragraph(intro_line)
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_intro.paragraph_format.line_spacing = 1.15
        p_intro.paragraph_format.space_after = Pt(12)
        
        # 3. Content
        for item in req.contenido:
            if item['tipo'] == 'parrafo':
                p = doc.add_paragraph(item['texto'])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(12)
            elif item['tipo'] == 'tabla':
                pass
        
        # 4. Signature Section (Deterministic)
        doc.add_paragraph("\n\n")
        sign_p = doc.add_paragraph("A T E N T A M E N T E\n\n\n__________________________________\n")
        sign_p.add_run(f"{req.representante_legal.upper()}\n").bold = True
        sign_p.add_run(f"{req.cargo_representante.upper()}\n")
        sign_p.add_run(f"{req.empresa_nombre.upper()}")
        sign_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to temp file
        fd, path = tempfile.mkstemp(suffix=".docx")
        try:
            doc.save(path)
            return FileResponse(path, filename=f"{req.titulo_documento.replace(' ', '_')}.docx")
        finally:
            os.close(fd)
            # Note: We should handle temp file deletion after response
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/docx/from-excel")
async def generate_from_excel(
    file: UploadFile = File(...),
    empresa_nombre: str = "EMPRESA",
    empresa_rfc: str = "",
    representante_legal: str = "",
    cargo_representante: str = "Representante Legal",
    convocante: str = "",
    licitacion_no: str = "",
    objeto: str = "",
    domicilio_fiscal: str = "",
    logo_path: str = "",
    workspace_id: str = "default"
):
    """Lee un Excel de cotización y genera el DOCUMENTO E2 - Propuesta Económica en Word"""
    import openpyxl
    from docx.shared import RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import logging
    
    logger = logging.getLogger("docx-gen")

    try:
        content = await file.read()
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        ws = wb.active
        logger.info(f"[Excel→E2] Leyendo hoja: '{ws.title}', dimensiones: {ws.dimensions}")

        # ============================================
        # PASO 1: INTERPRETAR LA HOJA DE CÁLCULO
        # ============================================
        rows = list(ws.iter_rows(values_only=True))
        
        # 1. Encontrar cabeceras
        headers = []
        header_row_idx = -1
        for i, row in enumerate(rows):
            cleaned = [str(c).strip() if c is not None else "" for c in row]
            non_empty = [c for c in cleaned if c and c != "None"]
            if len(non_empty) >= 4 and not headers:
                headers = cleaned
                header_row_idx = i
                break

        if not headers:
            return {"status": "error", "msg": "No se detectaron cabeceras válidas en el Excel."}

        # 2. Identificar índices de columnas clave
        desc_col = -1; price_col = -1; total_col = -1; qty_col = -1
        for idx, h in enumerate(headers):
            h_l = str(h).lower()
            if any(kw in h_l for kw in ['descrip', 'bien', 'concepto']): desc_col = idx
            if any(kw in h_l for kw in ['unitario', 'p.u.']): price_col = idx
            if any(kw in h_l for kw in ['acumulado', 'importe', 'monto total']) or (idx == len(headers)-1): total_col = idx
            if any(kw in h_l for kw in ['cant', 'qty']): qty_col = idx

        # 3. Agrupar partidas (manejo de filas múltiples por item)
        final_items = []
        current_item = None
        stop_keywords = ['total', 'iva', 'neto', 't.c.', 'pesos', 'dolares', 'subtotal', 'sub-total']

        for row in rows[header_row_idx+1:]:
            cleaned = [str(c).strip() if c is not None else "" for c in row]
            non_empty_cells = [c for c in cleaned if c and c != "None"]
            if not non_empty_cells: continue
            
            row_str = " ".join(cleaned).lower()
            if any(kw in row_str for kw in stop_keywords) and len(non_empty_cells) < 4:
                continue

            has_anchor = any(str(row[j]).strip() not in ["", "None"] for j in range(min(3, len(row))))
            
            val_p = str(row[price_col]).replace(",","").replace("$","").replace(" ", "").strip() if price_col < len(row) else ""
            val_t = str(row[total_col]).replace(",","").replace("$","").replace(" ", "").strip() if total_col < len(row) else ""
            try:
                p_num = float(val_p) if val_p else 0
                t_num = float(val_t) if val_t else 0
            except: p_num = t_num = 0

            if has_anchor or not final_items:
                current_item = list(cleaned)
                current_item[price_col] = p_num
                current_item[total_col] = t_num
                final_items.append(current_item)
            else:
                if desc_col >= 0 and row[desc_col] and str(row[desc_col]) != "None":
                    current_item[desc_col] = str(current_item[desc_col]) + "\n" + str(row[desc_col])
                if p_num > 0: current_item[price_col] = p_num
                if t_num > 0: current_item[total_col] = t_num

        data_rows = final_items
        subtotal = sum(item[total_col] if isinstance(item[total_col], (int, float)) else 0 for item in final_items)
        
        logger.info(f"[Excel→E2] Partidas reales extraídas: {len(data_rows)}")

        # ============================================
        # PASO 2: DETECTAR COLUMNAS DE PRECIO/CANTIDAD
        # ============================================
        price_keywords = ['precio', 'costo', 'importe', 'monto', 'valor', 'unit', 'p.u.', 'pu']
        qty_keywords = ['cantidad', 'cant', 'qty', 'piezas', 'pzas']
        
        price_cols = set()
        qty_cols = set()
        
        for i, h in enumerate(headers):
            h_l = h.lower()
            if any(kw in h_l for kw in price_keywords):
                price_cols.add(i)
            elif any(kw in h_l for kw in qty_keywords):
                qty_cols.add(i)
        
        # ============================================
        # PASO 3: CALCULAR TOTALES (SIN DUPLICAR)
        # ============================================
        # Buscamos la columna de "PRECIO ACUMULADO" o la última numérica
        total_col_idx = None
        for i in reversed(range(len(headers))):
            if i in price_cols or 'acumulado' in headers[i].lower() or 'total' in headers[i].lower():
                total_col_idx = i
                break
        
        subtotal = 0.0
        if total_col_idx is not None:
            for row in data_rows:
                val = str(row[total_col_idx]).replace(",", "").replace("$", "").replace(" ", "").strip()
                try:
                    subtotal += float(val)
                except:
                    pass
        
        iva_rate = 0.16
        iva_amount = subtotal * iva_rate
        total_con_iva = subtotal + iva_amount

        # ============================================
        # PASO 4: GENERAR DOCUMENTO WORD
        # ============================================
        def set_cell_shading(cell, color_hex):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), color_hex)
            cell._tc.get_or_add_tcPr().append(shading)

        def format_currency(val):
            try:
                num = float(str(val).replace(",", "").replace("$", "").replace(" ", "").strip())
                return f"${num:,.2f}"
            except:
                return str(val)

        # Compatibilidad con el renderizado de la tabla
        price_cols = {price_col, total_col}
        qty_cols = {qty_col}

        doc = Document()
        # ... (estilos y márgenes)
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        for section in doc.sections:
            section.top_margin = Inches(0.98); section.bottom_margin = Inches(0.98)
            section.left_margin = Inches(0.78); section.right_margin = Inches(0.78)

        if convocante:
            add_header(doc, convocante, licitacion_no, logo_path if logo_path and os.path.exists(logo_path) else None)

        doc.add_paragraph("\n")
        title = doc.add_heading("DOCUMENTO E2 – PROPUESTA ECONÓMICA", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datos empresa
        today = datetime.date.today().strftime("%d/%m/%Y")
        p_info = doc.add_paragraph(f"Fecha: {today}\nLicitante: {empresa_nombre.upper()}\nR.F.C.: {empresa_rfc}")
        p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p_info.runs: r.font.size = Pt(9)

        # Intro
        intro_txt = (f"En mi carácter de {cargo_representante} de la empresa {empresa_nombre.upper()}, "
                    f"con R.F.C. {empresa_rfc}, presento la propuesta económica para la licitación {licitacion_no}, "
                    f"relativa a {objeto or 'los trabajos descritos'}.")
        p_intro = doc.add_paragraph(intro_txt)
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # TABLA
        if headers and data_rows:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Header row
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h.upper()
                set_cell_shading(cell, '1a1a2e')
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)

            # Data rows
            for idx, row_data in enumerate(data_rows):
                row_cells = table.add_row().cells
                is_even = idx % 2 == 0
                for i, val in enumerate(row_data):
                    display_val = str(val) if val and str(val) != "None" else ""
                    
                    # Formatear: Solo columnas de PRECIO llevan $, las de CANTIDAD no.
                    if i in price_cols and display_val:
                        display_val = format_currency(display_val)
                    
                    row_cells[i].text = display_val
                    p = row_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if (i in price_cols or i in qty_cols) else WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs: r.font.size = Pt(8)
                    if is_even: set_cell_shading(row_cells[i], 'f0f4ff')

            # TOTALES FINAL
            row_sub = table.add_row().cells
            table.rows[-1].cells[0].merge(table.rows[-1].cells[-2])
            table.rows[-1].cells[0].text = "SUBTOTAL"
            table.rows[-1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.rows[-1].cells[-1].text = f"${subtotal:,.2f}"
            
            row_iva = table.add_row().cells
            table.rows[-1].cells[0].merge(table.rows[-1].cells[-2])
            table.rows[-1].cells[0].text = "I.V.A. (16%)"
            table.rows[-1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.rows[-1].cells[-1].text = f"${iva_amount:,.2f}"

            row_tot = table.add_row().cells
            table.rows[-1].cells[0].merge(table.rows[-1].cells[-2])
            table.rows[-1].cells[0].text = "TOTAL CON I.V.A."
            table.rows[-1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.rows[-1].cells[-1].text = f"${total_con_iva:,.2f}"
            for cell in table.rows[-1].cells:
                set_cell_shading(cell, '1a1a2e')
                for r in cell.paragraphs[0].runs: r.font.color.rgb = RGBColor(255, 255, 255); r.font.bold = True

        # Firma
        doc.add_paragraph("\n\n")
        sign = doc.add_paragraph(f"A T E N T A M E N T E\n\n\n__________________________________\n"
                                f"{representante_legal.upper()}\n{cargo_representante.upper()}\n{empresa_nombre.upper()}")
        sign.alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_footer(doc)
        output_path = os.path.join(f"/app/data/workspaces/{workspace_id}", "DOCUMENTO_E2_Presupuesto.docx")
        doc.save(output_path)
        return FileResponse(output_path, filename="DOCUMENTO_E2_Presupuesto.docx")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando Excel: {str(e)}")


@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "docx-gen"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8081)
